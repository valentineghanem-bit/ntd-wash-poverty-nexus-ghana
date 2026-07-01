"""
assemble_docx.py — Stage 10: assemble Q1 manuscript. Vancouver re-index by first appearance,
embed figures AND formatted tables, full front/back matter. Run qa/q1_gate.py after.
Manuscript .docx is a deliverable (NOT committed to git, Tenet 20).
"""
import re, os, sys
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import pandas as pd
from _labels import relabel_cols
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

MAN="manuscript/"; FIG="outputs/figures/"; TAB="outputs/tables/"

REFS = {
 "Hotez et al. 2009":"Hotez PJ, Fenwick A, Savioli L, Molyneux DH. Rescuing the bottom billion through control of neglected tropical diseases. Lancet. 2009;373(9674):1570–5.",
 "Branda et al. 2024":"Branda F, Cella E, Scarpa F, et al. Assessing the burden of neglected tropical diseases in low-income communities: challenges and solutions. Viruses. 2024;16(12):1932.",
 "Houweling et al. 2016":"Houweling TAJ, Karim-Kos HE, Kulik MC, et al. Socioeconomic inequalities in neglected tropical diseases: a systematic review. PLoS Negl Trop Dis. 2016;10(5):e0004546. doi:10.1371/journal.pntd.0004546.",
 "Magalhães et al. 2023":"Magalhães AR, Codeço CT, Svenning JC, et al. Neglected tropical diseases risk correlates with poverty and early ecosystem destruction. Infect Dis Poverty. 2023;12:32. doi:10.1186/s40249-023-01084-1.",
 "Boisson et al. 2016":"Boisson S, Engels D, Gordon BA, et al. Water, sanitation and hygiene for accelerating and sustaining progress on neglected tropical diseases: a new Global Strategy 2015–20. Int Health. 2016;8(Suppl 1):i19–i21.",
 "Naqvi et al. 2022":"Naqvi F, Cousens S, Awasthi S, et al. Interventions for neglected tropical diseases among children and adolescents: a meta-analysis. Pediatrics. 2022;149(Suppl 6):e2021053852G.",
 "WHO 2020":"World Health Organization. Ending the neglect to attain the Sustainable Development Goals: a road map for neglected tropical diseases 2021–2030. Geneva: WHO; 2020.",
 "WHO 2018":"World Health Organization Regional Office for Africa. Ghana eliminates trachoma, freeing millions from suffering and blindness. Brazzaville: WHO AFRO; 2018.",
 "WHO & UNICEF 2023":"WHO, UNICEF. Progress on household drinking water, sanitation and hygiene 2000–2022: special focus on gender. Geneva: WHO/UNICEF Joint Monitoring Programme; 2023.",
 "WHO-ESPEN 2026":"World Health Organization. Expanded Special Project for Elimination of Neglected Tropical Diseases (ESPEN): Ghana country profile. Brazzaville: WHO AFRO; 2026. https://espen.afro.who.int.",
 "Ahiadorme et al. 2026":"Ahiadorme M, Asori M, Manyeh AK. High-resolution mapping of onchocerciasis risk in Ghana using spatial machine learning models. BMC Infect Dis. 2026;26(1):1–15. doi:10.1186/s12879-026-13415-2.",
 "Nyerere & Mulwa 2026":"Nyerere N, Mulwa DF. Advanced machine learning approaches for predicting neglected tropical disease co-endemicity in Kenya. PLoS Negl Trop Dis. 2026;20(3):e0014156. doi:10.1371/journal.pntd.0014156.",
 "Kenu et al. 2014":"Kenu E, Ganu V, Calys-Tagoe BNL, et al. Application of geographical information system (GIS) technology in the control of Buruli ulcer in Ghana. BMC Public Health. 2014;14:724. doi:10.1186/1471-2458-14-724.",
 "Ablordey et al. 2015":"Ablordey AS, Vandelannoote K, Frimpong IA, et al. Whole genome comparisons suggest random distribution of Mycobacterium ulcerans genotypes in a Buruli ulcer endemic region of Ghana. PLoS Negl Trop Dis. 2015;9(3):e0003681. doi:10.1371/journal.pntd.0003681.",
 "Eneanya et al. 2024":"Eneanya OA, Delea MG, Cano J, et al. Predicting the environmental suitability and identifying climate and sociodemographic correlates of Guinea worm (Dracunculus medinensis) in Chad. Am J Trop Med Hyg. 2024;111(3_Suppl):26–35. doi:10.4269/ajtmh.23-0681.",
 "Adamou et al. 2026":"Adamou S, Laouali B, Garba N, et al. Niger: the first human onchocerciasis-endemic country in Africa to achieve elimination of infection transmission. Int Health. 2026. doi:10.1093/inthealth/ihag013.",
 "Vinkeles Melchers et al. 2024":"Vinkeles Melchers NVS, Agoro S, Togbey K, et al. Impact of ivermectin and vector control on onchocerciasis transmission in Togo. PLoS Negl Trop Dis. 2024;18(7):e0012312. doi:10.1371/journal.pntd.0012312.",
 "Nwane et al. 2025":"Bienvenu Nwane P, Nana-Djeunga HC, Toche NN, et al. Status of human onchocerciasis transmission in the Adamaoua region of Cameroon after 20 years of ivermectin mass distribution. PLoS Negl Trop Dis. 2025;19(3):e0011511. doi:10.1371/journal.pntd.0011511.",
 "Wolde et al. 2024":"Wolde HM, Getu M, Seid G, et al. Mapping the distribution of tuberculosis cases and associated factors in Central Ethiopia. BMC Public Health. 2024;24:2913. doi:10.1186/s12889-024-20343-w.",
 "Anselin 1995":"Anselin L. Local indicators of spatial association—LISA. Geogr Anal. 1995;27(2):93–115.",
 "Kelsall & Diggle 1995":"Kelsall JE, Diggle PJ. Non-parametric estimation of spatial variation in relative risk. Stat Med. 1995;14(21–22):2335–42.",
 "Openshaw 1984":"Openshaw S. The modifiable areal unit problem. CATMOG 38. Norwich: Geo Books; 1984.",
 "Moran 1950":"Moran PAP. Notes on continuous stochastic phenomena. Biometrika. 1950;37(1–2):17–23.",
 "Getis & Ord 1992":"Getis A, Ord JK. The analysis of spatial association by use of distance statistics. Geogr Anal. 1992;24(3):189–206.",
 "Tobler 1970":"Tobler WR. A computer movie simulating urban growth in the Detroit region. Econ Geogr. 1970;46(Suppl):234–40.",
 "Benjamini & Hochberg 1995":"Benjamini Y, Hochberg Y. Controlling the false discovery rate: a practical and powerful approach to multiple testing. J R Stat Soc B. 1995;57(1):289–300.",
 "Breiman et al. 1984":"Breiman L, Friedman JH, Olshen RA, Stone CJ. Classification and regression trees. Belmont (CA): Wadsworth; 1984.",
 "Breiman 2001":"Breiman L. Random forests. Mach Learn. 2001;45(1):5–32.",
 "Lundberg & Lee 2017":"Lundberg SM, Lee SI. A unified approach to interpreting model predictions. Adv Neural Inf Process Syst. 2017;30:4765–74.",
 "Pedregosa et al. 2011":"Pedregosa F, Varoquaux G, Gramfort A, et al. Scikit-learn: machine learning in Python. J Mach Learn Res. 2011;12:2825–30.",
 "Alkire & Foster 2011":"Alkire S, Foster J. Counting and multidimensional poverty measurement. J Public Econ. 2011;95(7–8):476–87.",
 "von Elm et al. 2007":"von Elm E, Altman DG, Egger M, et al. The Strengthening the Reporting of Observational Studies in Epidemiology (STROBE) statement. Ann Intern Med. 2007;147(8):573–7.",
 "Benchimol et al. 2015":"Benchimol EI, Smeeth L, Guttmann A, et al. The REporting of studies Conducted using Observational Routinely-collected health Data (RECORD) statement. PLoS Med. 2015;12(10):e1001885.",
 "Collins et al. 2015":"Collins GS, Reitsma JB, Altman DG, Moons KGM. Transparent reporting of a multivariable prediction model for individual prognosis or diagnosis (TRIPOD). Ann Intern Med. 2015;162(1):55–63.",
 "Collins et al. 2024":"Collins GS, Moons KGM, Dhiman P, et al. TRIPOD+AI statement: updated guidance for reporting clinical prediction models that use regression or machine learning methods. BMJ. 2024;385:e078378.",
 "Ghana Statistical Service 2021":"Ghana Statistical Service. Ghana 2021 Population and Housing Census: General Report. Accra: GSS; 2021.",
 "Ghana Statistical Service & ICF 2024":"Ghana Statistical Service, ICF. Ghana Demographic and Health Survey 2022. Accra and Rockville (MD): GSS and ICF; 2024.",
}

FIGCAP={
 2:("fig2_nwpri_choropleth.png","Figure 2. NTD WASH–Poverty Receptivity Index (NWPRI) across 261 Ghanaian districts (PCA composite, quintiles; colourblind-safe)."),
 3:("fig3_lisa_clusters.png","Figure 3. Local clusters of NTD receptivity (Local Moran's I on genuine-district deprivation, FDR-adjusted p < 0.05)."),
 4:("fig4_kde_poverty.png","Figure 4. Poverty-weighted kernel density surface of NTD receptivity over district centroids."),
 5:("fig5_ntd_overlay.png","Figure 5. Spatial overlay of receptivity hotspots, the documented NTD poverty-belt and southern Buruli foci."),
 6:("fig6_roc_calibration.png","Figure 6. ROC and calibration of receptivity classifiers under region-blocked spatial cross-validation."),
 7:("fig7_shap_summary.png","Figure 7. SHAP summary of determinants of NTD poverty-belt membership (random forest)."),
 8:("fig8_cart_tree.png","Figure 8. CART decision tree for NTD poverty-belt classification."),
}
TABLES={  # longest keys first for matching
 "4b":("table4b_ml_genuine_district_only.csv","Table 4b. Genuine-district-only model (no region-broadcast WASH; region-blocked CV).",None),
 "1":("table1_district_descriptives.csv","Table 1. Descriptive statistics of the district analytic frame (n = 261).",None),
 "2":("table2_global_morans_i.csv","Table 2. Global Moran's I for receptivity measures (queen contiguity, 999 permutations).",None),
 "3":("table3_lisa_HH_hotspots.csv","Table 3. High-High LISA receptivity hotspot districts (FDR < 0.05; top 15 of 31).",15),
 "4":("table4_ml_performance.csv","Table 4. Machine-learning performance predicting the NTD poverty-belt (region-blocked spatial CV).",None),
 "5":("table5_feature_importance.csv","Table 5. Random-forest feature importance (full model).",None),
 "6":("table6_robustness.csv","Table 6. Robustness checks: permutation null, south-only signal, anchor sensitivity.",None),
}

def load(f): return open(MAN+f,encoding="utf-8").read()
abstract=load("abstract.md"); im=load("draft_intro_methods.md")
res=load("draft_results.md"); dis=load("draft_discussion.md")
intro=im.split("## Methods")[0]; methods="## Methods"+im.split("## Methods")[1]
title=re.search(r"^# (.+)$",im,re.M).group(1)
body_order=[("Introduction",intro),("Methods",methods),("Results",res),("Discussion",dis)]
full_text=" ".join(t for _,t in body_order)

# ---- Vancouver re-index by first appearance (key-based, robust) -------------------------
pos={k:full_text.find("("+k) if ("("+k) in full_text else min([m.start() for m in re.finditer(re.escape(k),full_text)] or [10**9]) for k in REFS}
order=[k for k in sorted(REFS,key=lambda k:pos[k]) if pos[k]<10**9]
num={k:i+1 for i,k in enumerate(order)}
paren_re=re.compile(r"\(([^()]*\d{4}[^()]*)\)")
def vancouver(txt):
    def repl(m):
        segs=[s.strip() for s in m.group(1).split(";")]
        if segs and all(s in num for s in segs):
            return "["+",".join(str(num[s]) for s in segs)+"]"
        return m.group(0)
    return paren_re.sub(repl,txt)

# ---- docx helpers ----------------------------------------------------------------------
doc=Document(); st=doc.styles["Normal"]; st.font.name="Calibri"; st.font.size=Pt(11)
def H(t,l=1): return doc.add_heading(t,level=l)
def rich(p,para):
    for seg in re.split(r"(\*\*.*?\*\*)",para):
        r=p.add_run(re.sub(r"\*\*","",seg))
        if seg.startswith("**"): r.bold=True
def add_table(key):
    fn,cap,limit=TABLES[key]; path=TAB+fn
    if not os.path.exists(path): return
    df=pd.read_csv(path)
    if limit: df=df.head(limit)
    df=relabel_cols(df)
    cp=doc.add_paragraph(cap); cp.runs[0].italic=True; cp.runs[0].font.size=Pt(9)
    t=doc.add_table(rows=1,cols=len(df.columns)); t.style="Light List Accent 1"
    for j,c in enumerate(df.columns):
        cell=t.rows[0].cells[j]; cell.text=str(c)
        cell.paragraphs[0].runs[0].bold=True; cell.paragraphs[0].runs[0].font.size=Pt(8)
    for _,row in df.iterrows():
        cells=t.add_row().cells
        for j,v in enumerate(row):
            sval="" if (pd.isna(v) or str(v).lower()=="nan") else str(v)
            cells[j].text=sval
            if cells[j].paragraphs[0].runs: cells[j].paragraphs[0].runs[0].font.size=Pt(8)

# ---- title / front matter --------------------------------------------------------------
doc.add_heading(title,0)
doc.add_paragraph("Valentine Golden Ghanem").alignment=WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("Correspondence: valentineghanem@gmail.com").alignment=WD_ALIGN_PARAGRAPH.CENTER
rh=doc.add_paragraph(); rh.add_run("Running head: ").bold=True; rh.add_run("NTD WASH–poverty receptivity in Ghana")
# Abstract
H("Abstract",1)
for para in abstract.split("\n\n"):
    para=para.strip()
    if not para or para.startswith("# Abstract"): continue
    rich(doc.add_paragraph(),para.replace("**Keywords:**","Keywords:"))
ab=doc.add_paragraph(); ab.add_run("Abbreviations: ").bold=True
ab.add_run("AUC, area under the ROC curve; CART, classification and regression tree; CI, confidence interval; "
 "CRediT, Contributor Roles Taxonomy; DHS, Demographic and Health Survey; ESPEN, Expanded Special Project for "
 "Elimination of NTDs; FDR, false discovery rate; GHO, Global Health Observatory; GSS, Ghana Statistical Service; "
 "IQR, interquartile range; LISA, local indicators of spatial association; MAUP, modifiable areal unit problem; "
 "MDA, mass drug administration; MMDA, metropolitan/municipal/district assembly; NTD, neglected tropical disease; "
 "NWPRI, NTD WASH–Poverty Receptivity Index; PCA, principal component analysis; RF, random forest; "
 "SHAP, SHapley Additive exPlanations; WASH, water, sanitation and hygiene.")

# ---- body (figures + tables embedded at first citation) --------------------------------
fig_done=set(); tab_done=set()
for secname,txt in body_order:
    txt=vancouver(txt); H(secname,1)
    for block in txt.split("\n\n"):
        block=block.strip()
        if not block or block.startswith("# ") or block.startswith("*"): continue
        para=block
        if block.startswith("## "):
            lines=block.split("\n",1); sub=lines[0].lstrip("# ").strip()
            if sub.lower() not in (secname.lower(),"methods"): H(sub,2)
            if len(lines)<2 or not lines[1].strip(): continue
            para=lines[1].strip()
        rich(doc.add_paragraph(),para)
        for n,(fn,cap) in FIGCAP.items():
            if f"Figure {n}" in para and n not in fig_done and os.path.exists(FIG+fn):
                doc.add_picture(FIG+fn,width=Inches(5.4)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
                c=doc.add_paragraph(cap); c.runs[0].italic=True; c.runs[0].font.size=Pt(9); fig_done.add(n)
        for key in sorted(TABLES,key=len,reverse=True):
            if f"Table {key}" in para and key not in tab_done:
                add_table(key); tab_done.add(key)

# ---- declarations / back matter --------------------------------------------------------
H("Declarations",1)
def decl(h,b): p=doc.add_paragraph(); p.add_run(h+" ").bold=True; p.add_run(b)
decl("Ethics approval and consent.","Analysis of publicly available, de-identified aggregate secondary data (WHO-GHO/ESPEN, DHS Program, World Bank, Ghana Statistical Service); no human participants. Exempt from full ethical review under Ghana Health Service Ethics Review Committee guidance on secondary use of anonymised public data.")
decl("Data availability.","All source datasets are public (WHO Global Health Observatory and ESPEN, the DHS Program STATcompiler, the World Bank, and Ghana Statistical Service census tables). Derived master datasets, the district crosswalk, and all analysis code are available in the project repository and will be archived with a Zenodo DOI on publication.")
decl("Code availability.","Python pipeline (build_master, build_ntd_anchor, spatial_analysis, ml_receptivity, robustness_checks, assemble) released in the repository; reproducible end-to-end.")
decl("Funding.","This research received no specific grant from any funding agency in the public, commercial, or not-for-profit sectors.")
decl("Competing interests.","The author declares no competing interests.")
decl("Author contributions (CRediT).","Valentine Golden Ghanem: conceptualization, methodology, software, formal analysis, data curation, visualization, writing — original draft, writing — review & editing.")
decl("Acknowledgements.","The author thanks the WHO-ESPEN, the DHS Program, and the Ghana Statistical Service for open data access.")
decl("Reporting.","The study is reported per STROBE and RECORD; the prediction model follows TRIPOD+AI. Completed checklists are provided as Supplementary Files S1–S3.")
decl("Supplementary material.","S1 STROBE checklist; S2 RECORD checklist; S3 TRIPOD+AI checklist; S4 supplementary figures (kernel density, ROC/calibration, CART) and tables (full LISA, robustness).")

# ---- references ------------------------------------------------------------------------
H("References",1)
for k in order: doc.add_paragraph(f"{num[k]}. {REFS[k]}")

out=MAN+"NTD_WASH_Poverty_Ghana_manuscript.docx"; doc.save(out)
print("Saved:",out)
print(f"Refs: {len(order)} | figures: {sorted(fig_done)} | tables: {sorted(tab_done)}")
uncited=[k for k in REFS if k not in num]
if uncited: print("WARNING uncited refs (dropped):",uncited)
