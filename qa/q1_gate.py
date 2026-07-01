"""
q1_gate.py — DETERMINISTIC Q1 manuscript gate. Run BEFORE declaring production done.
Hard-fails (exit 1) on any unmet item in _system/Q1_DELIVERABLE_STANDARDS.md §1.
Purpose: make the Q1 standard a failing test, not a thing to remember.
"""
import sys, os, re
from docx import Document

DOCX="manuscript/NTD_WASH_Poverty_Ghana_manuscript.docx"
checks=[]
def chk(name, ok, detail=""):
    checks.append((name, bool(ok), detail))

if not os.path.exists(DOCX):
    print("FAIL: manuscript .docx not built"); sys.exit(1)
d=Document(DOCX)
text="\n".join(p.text for p in d.paragraphs)
low=text.lower()
n_tables=len(d.tables)
n_figs=sum(1 for _ in d.inline_shapes)
n_words=len(text.split())

# reference count = lines like "12. Author..."
refs=re.findall(r"\n\s*\d+\.\s+[A-Z]", "\n"+text)
n_refs=len(refs)

# abstract words = structured abstract body only (stop at Keywords/Abbreviations/Introduction)
abm=re.search(r"Abstract(.*?)(?:Keywords|Abbreviations|Introduction)", text, re.S)
abs_words=len(abm.group(1).split()) if abm else 0

chk("Tables embedded as formatted tables (>=6)", n_tables>=6, f"{n_tables} tables")
chk("Figures embedded (>=4, <=8 main)", 4<=n_figs<=8, f"{n_figs} figures")
chk("References at Q1 depth (>=35)", n_refs>=35, f"{n_refs} refs")
chk("Structured abstract <=300 words", 0<abs_words<=300, f"{abs_words} words")
chk("Manuscript length 3000-6000 words", 3000<=n_words<=6000, f"{n_words} words")
chk("Data availability statement", "data availability" in low)
chk("Funding statement", "funding" in low)
chk("Competing interests / COI", ("competing interest" in low) or ("conflict of interest" in low))
chk("Author contributions (CRediT)", ("author contribution" in low) or ("credit" in low))
chk("Running head", "running head" in low)
chk("Abbreviations list", "abbreviation" in low)
chk("Ethics/IRB statement", ("ethic" in low) or ("irb" in low) or ("review committee" in low))
chk("STROBE reporting", "strobe" in low)
chk("RECORD reporting", "record" in low)
chk("TRIPOD-AI (ML reporting)", "tripod" in low)
chk("Keywords present", "keyword" in low)
chk("Limitations explicit", "limitation" in low)
chk("Vancouver numbered in-text [n]", bool(re.search(r"\[\d+\]", text)))

# QUALITY: no raw machine variable names (snake_case / _rgnbc / _pca) in embedded tables
raw_re=re.compile(r"[a-z]+_[a-z]+|_rgnbc|_pca|pooledOOF|bal_acc")
table_raw=[]
for ti,tb in enumerate(d.tables):
    for row in tb.rows:
        for c in row.cells:
            if raw_re.search(c.text): table_raw.append((ti,c.text)); break
chk("Tables use readable labels (no snake_case)", len(table_raw)==0,
    f"{len(table_raw)} raw cells e.g. {table_raw[0][1] if table_raw else ''}")

print("="*64); print("Q1 MANUSCRIPT GATE"); print("="*64)
fails=0
for name,ok,detail in checks:
    print(f"  [{'PASS' if ok else 'FAIL'}] {name}" + (f"  ({detail})" if detail else ""))
    fails += (not ok)
print("="*64)
print(f"{len(checks)-fails}/{len(checks)} passed | {fails} FAILING")
sys.exit(1 if fails else 0)
